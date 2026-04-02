import re
import uuid
from ..core.logger import get_logger
from ..core.ai_clients import call_llm

logger = get_logger(__name__)

class NLPService:
    def __init__(self, model_name='all-MiniLM-L6-v2'):
        logger.info("Initializing NLPService (Lazy Loading with LLM)...")
        self.model_name = model_name

    def calculate_similarity(self, text1, text2):
        """
        Calculates similarity using LLM evaluation.
        Returns a score between 0 and 1.
        """
        if not text1 or not text2:
            return 0.0
        
        from .interview import evaluate_answer_content
        # Note: evaluate_answer_content returns a dict with "score"
        result = evaluate_answer_content(question=text2, answer=text1)
        return float(result.get("score", 0.0)) / 10.0 # Normalize 0-10 to 0-1

    def extract_text_from_file(self, file_path):
        """Extracts raw text from .txt, .pdf, or .docx files."""
        ext = os.path.splitext(file_path)[1].lower()
        content = ""

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif ext == '.pdf':
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                content = "\n".join([page.get_text() for page in doc])
                doc.close()
            elif ext == '.docx':
                from docx import Document
                doc = Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                raise ValueError(f"Unsupported file format for text extraction: {ext}")
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            return ""

        # Clean text
        content = "".join(char for char in content if char.isprintable() or char in "\n\r\t")
        return content.strip()

    def arrange_resume_with_ai(self, resume_text: str) -> str:
        """Uses AI to summarize and structure raw resume text for better prompt generation."""
        if not resume_text:
            return ""

        system_prompt = (
            "You are a Professional Resume Analyzer. Your task is to extract and structure "
            "raw resume text into a clean professional summary optimized for generating interview questions."
        )
        
        prompt = (
            "Extract the following details from the raw resume text below. If a detail is missing, omit it.\n\n"
            "Format the output strictly as follows:\n"
            "- Core Tech Stack: [List primary languages, frameworks, and tools]\n"
            "- Total Experience: [X years/months]\n"
            "- Seniority Level: [Junior/Mid/Senior/Lead/etc.]\n"
            "- Key Achievements: [Bullet points of top 3 projects or achievements]\n"
            "- Professional Summary: [Max 3 sentences summarizing their background]\n\n"
            f"Raw Resume Text:\n{resume_text[:8000]}" # Truncate to save tokens for extraction
        )

        try:
            logger.info("Calling LLM to arrange resume text...")
            structured_summary = call_llm(prompt, system_prompt=system_prompt)
            return structured_summary or resume_text # Fallback to raw if AI fails
        except Exception as e:
            logger.error(f"Error arranging resume with AI: {e}")
            return resume_text

    def extract_qa_from_file(self, file_path, questions_only=False):
        """Extracts Q&A pairs (or just questions) from .txt, .pdf, .docx, or .xlsx files."""
        ext = os.path.splitext(file_path)[1].lower()
        content = ""

        try:
            if ext in ['.xlsx', '.xls']:
                import pandas as pd
                # For Excel, we use pandas
                df = pd.read_excel(file_path)
                
                # Priority 1: Search for common column names
                target_col = None
                common_names = ['question', 'questions', 'q']
                for col in df.columns:
                    if str(col).lower() in common_names:
                        target_col = col
                        break
                
                if target_col is not None:
                    questions = df[target_col].dropna().astype(str).tolist()
                elif len(df.columns) >= 2:
                    # Priority 2: 2nd column (skip serial no)
                    questions = df.iloc[:, 1].dropna().astype(str).tolist()
                else:
                    # Priority 3: 1st column
                    questions = df.iloc[:, 0].dropna().astype(str).tolist()
                
                return [{'question': q.strip()} for q in questions if q.strip()]

            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif ext == '.pdf':
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                content = "\n".join([page.get_text() for page in doc])
                doc.close()
            elif ext == '.docx':
                from docx import Document
                doc = Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            return []

        # Debug: Print first 500 chars to see what was extracted
        logger.debug(f"--- DOCUMENT CONTENT PREVIEW ({ext}) ---")
        logger.debug(content[:500])
        logger.debug("---------------------------------------")

        return self.parse_qa_pairs(content, questions_only=questions_only)

    def parse_qa_pairs(self, text, questions_only=False):
        """
        Regex-based parser to find Question and Answer patterns.
        If questions_only is True, it extracts everything as questions,
        splitting by patterns like "Q:", "1.", etc.
        """
        # Clean text
        text = "".join(char for char in text if char.isprintable() or char in "\n\r\t")
        
        # Patterns for questions and answers
        # Simplified patterns to reduce redundant alternatives (fixes lint)
        q_pattern = re.compile(r'^\s*(?:Q|Question|Ques|Q\s*):?\d*[\.\)]?\s*(.*)', re.IGNORECASE)
        a_pattern = re.compile(r'^\s*(?:A|Answer|Ans|Reference|R|Ref):?\s*(.*)', re.IGNORECASE)

        lines = text.split('\n')
        qa_pairs = []
        current_q = None
        current_a = None

        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            q_match = q_pattern.match(line)
            if q_match:
                # Save previous pair
                if current_q:
                    if questions_only:
                        qa_pairs.append({'question': current_q.strip()})
                    elif current_a:
                        qa_pairs.append({'question': current_q.strip(), 'answer': current_a.strip()})
                
                current_q = q_match.group(1)
                current_a = None
                continue

            if not questions_only:
                a_match = a_pattern.match(line)
                if a_match:
                    current_a = a_match.group(1)
                    continue

            # Continuation of previous line
            if current_a is not None:
                current_a += " " + line
            elif current_q is not None:
                current_q += " " + line
            elif questions_only and not current_q:
                current_q = line

        # Add the final pair
        if current_q:
            if questions_only:
                qa_pairs.append({'question': current_q.strip()})
            elif current_a:
                qa_pairs.append({'question': current_q.strip(), 'answer': current_a.strip()})

        logger.info(f"Extraction result: Found {len(qa_pairs)} pairs.")
        return qa_pairs
